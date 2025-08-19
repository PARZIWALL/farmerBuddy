import io
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import streamlit as st
from docx import Document
import re
from datetime import datetime

# ---------------- CONFIG ----------------
SCHEMES_INFO = {
    "PM-Kisan": {
        "name": "Pradhan Mantri Kisan Samman Nidhi Yojana",
        "description": "Direct income support of ₹6,000 per year for land-holding farmer families.",
        "eligibility": "All land-holding farmer families, subject to certain exclusion criteria (e.g., income tax payers, high-pensioners).",
        "documents_required": [
            "Land ownership documents (e.g., Khasra/Khatauni)",
            "Aadhaar Card (Mandatory for eKYC)",
            "Bank Account Passbook (for direct benefit transfer)",
            "Proof of Citizenship",
            "Passport Size Photograph"
        ],
        "submission_notes": "Submit the filled form and all documents to a Common Service Center (CSC), your local Patwari, or a designated Nodal Officer. Online self-registration is also available on the official PM-KISAN portal.",
        "template_file": "templates/pm-kisan_application_form.docx",
        "official_link": "https://pmkisan.gov.in/"
    },
    "PM-KMY": {
        "name": "Pradhan Mantri Kisan Maan Dhan Yojana",
        "description": "A voluntary and contributory pension scheme for small and marginal farmers.",
        "eligibility": "Small and Marginal Farmers (SMF) aged between 18 and 40 with cultivable land up to 2 hectares.",
        "documents_required": [
            "Aadhaar Card",
            "Bank Passbook or Bank Account Details",
            "Land Records (Khasra/Khatauni)",
            "Passport Size Photograph"
        ],
        "submission_notes": "Enrollment is done at the nearest Common Service Center (CSC). You will need to provide the filled form and documents, and the CSC will complete the online registration for you. PM-KISAN beneficiaries can opt for auto-debit of their contributions from the benefits they receive.",
        "template_file": "templates/pm-kmy_application_form.docx",
        "official_link": "https://pmkmy.gov.in/"
    },
    "KCC": {
        "name": "Kisan Credit Card",
        "description": "Concessional credit facility for agricultural and allied activities, offering short-term loans up to ₹3 lakh.",
        "eligibility": "All farmers, including owner cultivators, tenant farmers, sharecroppers, and members of Self Help Groups (SHGs) or Joint Liability Groups (JLGs).",
        "documents_required": [
            "Duly filled and signed application form",
            "Aadhaar Card",
            "PAN Card",
            "Land ownership/cultivation documents (e.g., land records, lease agreement)",
            "Bank Account Passbook",
            "Passport size photographs (2-3 copies)",
            "Income Certificate",
            "Any other security documents required by the bank"
        ],
        "submission_notes": "Submit the completed application form and all required documents to the bank's branch where you wish to open the account. The bank will then verify the details and process your application.",
        "template_file": "templates/kcc_application_form.docx",
        "official_link": "https://pmkisan.gov.in/Documents/Kcc.pdf"
    },
    "PMFBY": {
        "name": "Pradhan Mantri Fasal Bima Yojana",
        "description": "A crop insurance scheme that provides financial support to farmers suffering crop loss/damage arising out of unforeseen events.",
        "eligibility": "All farmers, including sharecroppers and tenant farmers, growing notified crops in a notified area.",
        "documents_required": [
            "Aadhaar Card",
            "Bank Account Passbook",
            "Land Records (Khatauni/Patwari records)",
            "Sowing Certificate from Village Revenue Officer",
            "Passport Size Photograph"
        ],
        "submission_notes": "The form must be submitted to the nearest bank branch, cooperative society, or CSC. Farmers who have taken institutional loans will have their premiums deducted automatically.",
        "template_file": "templates/pmfby_application_form.docx",
        "official_link": "https://pmfby.gov.in/"
    },
    "Soil Health Card": {
        "name": "Soil Health Card Scheme",
        "description": "A scheme to provide a Soil Health Card to all farmers, containing information on nutrient status and recommendations for soil improvement.",
        "eligibility": "All farmers with agricultural land.",
        "documents_required": [
            "Aadhaar Card",
            "Land records (Khatauni)",
            "Mobile number",
            "Copy of Bank Passbook"
        ],
        "submission_notes": "The form should be submitted to the local Agriculture Department or Krishi Vigyan Kendra. A soil sample will be collected from your farm for testing, and the card will be issued based on the results.",
        "template_file": "templates/soil_health_card_form.docx",
        "official_link": "https://soilhealth.dac.gov.in/"
    },
    "PM-KISAN SAMMAN": {
        "name": "Pradhan Mantri Kisan Samman Nidhi - Samman Patra (Certificate)",
        "description": "A scheme providing certificates to all PM-KISAN beneficiaries, verifying their enrollment in the scheme.",
        "eligibility": "All enrolled beneficiaries of the PM-KISAN scheme.",
        "documents_required": [
            "Aadhaar Card",
            "PM-KISAN registration number",
            "Land records (Khatauni)"
        ],
        "submission_notes": "The certificate can be downloaded online from the PM-KISAN official website or obtained from a CSC. The form is for verifying and updating beneficiary details before the certificate is issued.",
        "template_file": "templates/pm_kisan_samman_patra_form.docx",
        "official_link": "https://pmkisan.gov.in/"
    }
}

# ---------------- CORE DYNAMIC FUNCTIONS ----------------

def guess_input_type(label_text: str) -> str:
    """Guesses the input type based on keywords in the label."""
    text = label_text.lower()
    if any(keyword in text for keyword in ["date", "दिनांक", "तारीख"]):
        return "date"
    if any(keyword in text for keyword in ["mobile", "phone", "मोबाइल", "फोन"]):
        return "phone"
    if any(keyword in text for keyword in ["aadhaar", "uid", "आधार", "pan"]):
        return "text"
    if any(keyword in text for keyword in ["photo", "फोटो", "photograph"]):
        return "photo"
    if any(keyword in text for keyword in ["area", "क्षेत्र", "acre", "hectare", "amount"]):
        return "number"
    if any(keyword in text for keyword in ["account", "खाता", "ifsc", "name of bank"]):
        return "text"
    return "text"

def analyze_docx_dynamically(docx_file_path: str) -> List[Dict[str, Any]]:
    """
    Dynamically analyzes the DOCX file to find potential fillable fields.
    Prioritizes tables and then looks for standalone labels.
    """
    st.info("🧠 Dynamically analyzing DOCX structure...")
    detected_fields = []
    
    try:
        doc = Document(docx_file_path)
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return []

    # 1. Analyze Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    # Look for an empty cell in the same row, to the right, to serve as the fillable field
                    if col_idx + 1 < len(row.cells) and not row.cells[col_idx + 1].text.strip():
                        label_text = cell_text
                        
                        # Check for duplicates
                        if not any(f.get("question") == label_text and f.get("type") == "table" for f in detected_fields):
                            detected_fields.append({
                                "type": "table",
                                "question": label_text,
                                "input_type": guess_input_type(label_text),
                                "table_idx": table_idx,
                                "row_idx": row_idx,
                                "col_idx": col_idx + 1
                            })
    
    # 2. Analyze Paragraphs
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        # Regex to find a label ending with a colon, dots, or underscores
        match = re.match(r"(.*?)(:|\.|\_)\s*$", text)
        if match and not text.endswith((".docx", ".pdf", "doc")) and len(match.group(1).strip()) > 3:
            label_text = match.group(1).strip()
            
            # Ensure this isn't a duplicate of a table field
            if not any(f.get("question") == label_text for f in detected_fields):
                detected_fields.append({
                    "type": "paragraph",
                    "question": label_text,
                    "input_type": guess_input_type(label_text),
                    "para_idx": para_idx
                })

    if not detected_fields:
        st.warning("⚠️ No fillable fields detected automatically. The form may be non-standard.")
        return []
        
    st.success(f"✅ Found {len(detected_fields)} potential fillable fields!")
    return detected_fields

def fill_docx_template(docx_file_path: str, fields: List[Dict[str, Any]], photo_files: Dict = None) -> Optional[bytes]:
    """
    Fills a DOCX template with user answers.
    """
    try:
        doc = Document(docx_file_path)
        
        # 1. Fill fields in tables
        for field in [f for f in fields if f["type"] == "table"]:
            answer = field.get("answer", "")
            if answer:
                doc.tables[field["table_idx"]].rows[field["row_idx"]].cells[field["col_idx"]].text = str(answer)

        # 2. Fill fields in paragraphs
        for field in [f for f in fields if f["type"] == "paragraph"]:
            answer = field.get("answer", "")
            if answer:
                paragraph = doc.paragraphs[field["para_idx"]]
                # Ensure a space is added to separate the label from the answer
                paragraph.add_run(f" {str(answer)}")
            
        if photo_files:
            st.warning("Photo uploads are not currently supported for DOCX templates and will be ignored.")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except Exception as e:
        st.error(f"Error filling DOCX file: {e}")
        return None

# ---------------- STREAMLIT UI ----------------
def main():
    st.set_page_config(page_title="Krishi Sahayak - Smart Farmer Application System", layout="wide", initial_sidebar_state="expanded")
    
    st.markdown("""
        <style>
        .main-header { background: linear-gradient(90deg, #2E7D32, #4CAF50); padding: 1rem; border-radius: 10px; color: white; text-align: center; margin-bottom: 2rem; }
        .scheme-card { background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 8px; padding: 1rem; margin: 0.5rem 0;}
        .success-box { background: #d4edda; border: 1px solid #c3e6cb; color: #155724; padding: 1rem; border-radius: 5px; margin: 1rem 0;}
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div class="main-header">
            <h1>🌾 Krishi Sahayak - Smart Farmer Application System</h1>
            <p>A dynamic, AI-powered solution for Indian Government Farmer Schemes</p>
        </div>
    """, unsafe_allow_html=True)
    
    with st.sidebar:
        st.header("📋 Select Scheme")
        scheme_options = list(SCHEMES_INFO.keys())
        selected_scheme = st.selectbox("Choose Government Scheme:", scheme_options)
        
        if selected_scheme:
            scheme_info = SCHEMES_INFO[selected_scheme]
            st.markdown(f"""
                <div class="scheme-card">
                    <h4>{scheme_info['name']}</h4>
                    <p><strong>Description:</strong> {scheme_info['description']}</p>
                    <p><strong>Eligibility:</strong> {scheme_info['eligibility']}</p>
                </div>
            """, unsafe_allow_html=True)
            
            st.markdown("### 📄 Required Documents:")
            for doc in scheme_info['documents_required']:
                st.markdown(f"✓ {doc}")
            st.markdown(f"[🔗 Official Website]({scheme_info['official_link']})")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📤 Upload Application Form")
        
        uploaded_file = st.file_uploader(
            "Upload DOCX Application Form", 
            type=["docx"],
            help="Upload the official application form in .docx format to be filled"
        )
        
        template_file = None
        if uploaded_file:
            template_path = Path("templates") / uploaded_file.name
            template_path.parent.mkdir(exist_ok=True)
            with open(template_path, "wb") as f:
                f.write(uploaded_file.read())
            template_file = str(template_path)
            st.success("✅ Custom form uploaded successfully!")
        
        if template_file:
            st.info("""
            💡 **Prompt for Correct Placement:**
            - For best results, use **tables** for structured data (e.g., Name | [fillable cell]).
            - For standalone fields, end the label with a **colon (:)** or a series of **dots (...)**.
            - This system will automatically detect these patterns and fill the adjacent space.
            """)
            if st.button("🧠 Analyze Form to Find Fields", type="primary"):
                with st.spinner("Scanning DOCX to find fillable fields..."):
                    fields = analyze_docx_dynamically(template_file)
                    st.session_state["fields"] = fields
                    st.session_state["template_file"] = template_file
                    st.session_state["scheme_type"] = selected_scheme
            
            if "fields" in st.session_state and st.session_state["fields"]:
                with st.expander("🔍 Detected Fields Preview"):
                    for i, field in enumerate(st.session_state["fields"][:5]):
                        st.write(f"{i+1}. **{field['question']}** (type: {field['input_type']}, source: {field['type']})")
    
    with col2:
        st.header("ℹ️ Quick Info")
        st.info("""
            **How it works:**
            1. Upload a DOCX form template.
            2. The AI scans the document for tables and special characters to find fillable fields.
            3. Fill in the details for each field in the form below.
            4. Download your completed application as a new DOCX file.
        """)
    
    if "fields" in st.session_state and st.session_state["fields"]:
        st.header("📝 Fill Application Details")
        
        answers = {}
        
        with st.form("farmer_application_form"):
            for field in st.session_state["fields"]:
                unique_key = f"field_{field['type']}_{field.get('table_idx', '')}_{field.get('row_idx', '')}_{field.get('col_idx', '')}_{field.get('para_idx', '')}"
                
                # Create the appropriate Streamlit widget based on the guessed input type
                if field["input_type"] == "photo":
                    st.write(f"**Photo field detected, but DOCX photo filling is not supported.**")
                    answers[unique_key] = None
                elif field["input_type"] == "date":
                    answers[unique_key] = st.date_input(field["question"], key=unique_key)
                elif field["input_type"] == "number":
                    answers[unique_key] = st.number_input(field["question"], min_value=0.0, step=0.01, key=unique_key)
                else:
                    answers[unique_key] = st.text_input(field["question"], key=unique_key)
            
            submitted = st.form_submit_button("🚀 Generate Filled Application", type="primary")
        
        if submitted:
            with st.spinner("Generating your filled application..."):
                for field in st.session_state["fields"]:
                    unique_key = f"field_{field['type']}_{field.get('table_idx', '')}_{field.get('row_idx', '')}_{field.get('col_idx', '')}_{field.get('para_idx', '')}"
                    answer_value = answers.get(unique_key)
                    
                    if field["input_type"] == "date" and answer_value:
                        field["answer"] = answer_value.strftime("%d/%m/%Y")
                    elif field["input_type"] == "photo":
                        field["answer"] = ""
                    else:
                        field["answer"] = str(answer_value) if answer_value is not None else ""
                
                # Note: photo_files is not used for DOCX, so we pass an empty dict
                final_docx_bytes = fill_docx_template(st.session_state["template_file"], st.session_state["fields"], {})
                
                if final_docx_bytes:
                    st.markdown("""<div class="success-box"><h4>🎉 Application Successfully Generated!</h4><p>Your application has been filled and is ready for download.</p></div>""", unsafe_allow_html=True)
                    
                    filename = f"{st.session_state['scheme_type']}_application_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    st.download_button(label="📥 Download Filled Application DOCX", data=final_docx_bytes, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary")
                    
                    st.header("📋 Next Steps")
                    scheme_info = SCHEMES_INFO.get(st.session_state["scheme_type"], {})
                    st.success(f"""
                    **What to do next:**
                    1. 📄 Print the downloaded application form
                    2. ✍️ Sign the form where required
                    3. 📎 Attach all required documents (see sidebar)
                    4. 🏛️ Submit to nearest CSC/Bank/Government office
                    """)
                else:
                    st.error("❌ Failed to generate the final DOCX file. Please check your inputs.")

if __name__ == "__main__":
    Path("templates").mkdir(exist_ok=True)
    main()