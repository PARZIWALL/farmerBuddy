# app.py
import os
import json
import time
import re
import uuid
import sqlite3
import hashlib
from datetime import datetime
from flask import Flask, render_template, request, send_from_directory, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from docx import Document
import google.generativeai as genai
import pytesseract
import fitz  # PyMuPDF
from PIL import Image
from flask_cors import CORS



# -------------------------
# Configuration
# -------------------------
UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'generated_forms'
TEMPLATE_FOLDER = 'application_templates'
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'docx'}
DB_PATH = 'data.db'

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GENERATED_FOLDER'] = GENERATED_FOLDER
app.config['TEMPLATE_FOLDER'] = TEMPLATE_FOLDER
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your_super_secret_key')  # change for production

# Gemini config (optional)
GEMINI_API_KEY = ""
modelname= "gemini-2.0-flash"
if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(modelname)
    except Exception as e:
        print(f"Error configuring Gemini API: {e}")
else:
    print("GEMINI_API_KEY not set — RAG features will be unavailable.")

# -------------------------
# Example schemes
# -------------------------
SCHEMES = {
    'pm-kisan': {
        'name': 'Pradhan Mantri Kisan Samman Nidhi (PM-Kisan)',
        'template_file': 'pm-kisan_new_application_form_english.docx',
        'required_documents': [
            'Aadhaar Card',
            'Bank Passbook/Statement',
            'Land Records (Khasra/Khatauni)'
        ],
        'additional_documents': [
            'Aadhaar Card photocopy',
            'Bank Passbook photocopy',
            'Land ownership documents',
            '2 Passport size photographs',
        ]
    },
    'kcc': {
        'name': 'Kisan Credit Card (KCC)',
        'template_file': 'kcc_application_format.docx',
        'required_documents': [
            'Aadhaar Card',
            'PAN Card',
            'Bank Passbook',
            'Land Records',
        ],
        'additional_documents': [
            'Aadhaar Card photocopy',
            'PAN Card photocopy',
            'Land ownership proof',
            '2 Passport size photographs'
        ]
    },
    'ridf': {
        'name': 'Rural Infrastructure Development Fund (RIDF)',
        'template_file': 'RIDF G.APPLICATION FORM (1).docx',
        'required_documents': [
            'Project Proposal',
            'Land Documents',
            'Aadhaar Card of authorized person',
            'Bank Passbook'
        ],
        'additional_documents': [
            'Detailed project proposal',
            'Land ownership documents',
            'Identity and address proof of promoters',
            'Bank account statements'
        ]
    }
}

# -------------------------
# DB helpers
# -------------------------
def get_db_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db_conn()
    cur = conn.cursor()
    # users table stores internal uuid and optional hashed aadhaar
    cur.execute("""
    CREATE TABLE IF NOT EXISTS users (
        user_id TEXT PRIMARY KEY,
        aadhaar_hash TEXT,
        created_at TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT,
        filename TEXT,
        scheme_id TEXT,
        doc_type TEXT,
        text TEXT,
        metadata TEXT,
        chunk_index INTEGER,
        created_at TEXT
    )
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_documents_user ON documents(user_id)")
    conn.commit()
    conn.close()

def save_user_if_new(user_id, aadhaar_hash=None):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT user_id FROM users WHERE user_id = ?", (user_id,))
    if not cur.fetchone():
        cur.execute("INSERT INTO users (user_id, aadhaar_hash, created_at) VALUES (?, ?, ?)",
                    (user_id, aadhaar_hash, datetime.utcnow().isoformat()))
        conn.commit()
    else:
        # If user exists but aadhaar_hash provided and not present, update it
        if aadhaar_hash:
            cur.execute("SELECT aadhaar_hash FROM users WHERE user_id = ?", (user_id,))
            current = cur.fetchone()
            # If existing record has no aadhaar_hash, set it
            if current and not current['aadhaar_hash']:
                cur.execute("UPDATE users SET aadhaar_hash = ? WHERE user_id = ?", (aadhaar_hash, user_id))
                conn.commit()
    conn.close()

def save_document_record(user_id, filename, scheme_id, text, doc_type=None, metadata=None, chunk_index=-1):
    conn = get_db_conn()
    cur = conn.cursor()
    meta_json = json.dumps(metadata or {})
    cur.execute("""
        INSERT INTO documents (user_id, filename, scheme_id, doc_type, text, metadata, chunk_index, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (user_id, filename, scheme_id, doc_type or '', text, meta_json, chunk_index, datetime.utcnow().isoformat()))
    conn.commit()
    conn.close()

def get_documents_by_user(user_id):
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM documents WHERE user_id = ? ORDER BY created_at ASC", (user_id,))
    rows = cur.fetchall()
    conn.close()
    return [dict(r) for r in rows]

# -------------------------
# Utilities & OCR helpers
# -------------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_field_name_from_label(label):
    name = label.lower().strip().replace(':', '').replace('*', '').replace('.', '')
    name = name.replace(' ', '_').replace('-', '_')
    name = "_".join(filter(None, name.split('_')))
    return name

def find_aadhaar_in_text(text):
    """Returns first 12-digit sequence if found (naive)."""
    m = re.search(r'\b(\d{12})\b', text)
    return m.group(1) if m else None

def chunk_text_simple(text, words_per_chunk=400):
    words = text.split()
    if not words:
        return []
    chunks = []
    for i in range(0, len(words), words_per_chunk):
        chunks.append(" ".join(words[i:i+words_per_chunk]))
    return chunks

def ocr_image_with_confidence(pil_image, lang='eng'):
    """
    Returns dict: {'text': str, 'avg_confidence': float, 'words': [...], 'confs': [...]}
    """
    try:
        img = pil_image.convert('RGB')
        data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
        words = []
        confs = []
        for i, w in enumerate(data.get('text', [])):
            if w and w.strip():
                words.append(w)
                try:
                    conf = float(data['conf'][i])
                except Exception:
                    conf = 0.0
                confs.append(conf)
        full_text = " ".join(words)
        avg_conf = float(sum(confs) / len(confs)) if confs else 0.0
        return {'text': full_text, 'avg_confidence': avg_conf, 'words': words, 'confs': confs}
    except Exception as e:
        print("ocr_image_with_confidence error:", e)
        return {'text': '', 'avg_confidence': 0.0, 'words': [], 'confs': []}

def extract_text_from_file(filepath):
    """
    Returns a dict:
      {
        'text': '...',           # extracted text (string)
        'source': 'pdf|image|docx',
        'avg_conf': float_or_None
      }
    """
    text = ""
    source = None
    avg_conf = None
    try:
        ext = filepath.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            source = 'pdf'
            with fitz.open(filepath) as doc:
                pages_text = []
                for page in doc:
                    pages_text.append(page.get_text())
                text = "\n\n".join(pages_text)
        elif ext in ('png', 'jpg', 'jpeg'):
            source = 'image'
            img = Image.open(filepath)
            ocr_res = ocr_image_with_confidence(img)
            text = ocr_res['text']
            avg_conf = ocr_res['avg_confidence']
        elif ext == 'docx':
            source = 'docx'
            doc = Document(filepath)
            parts = []
            for para in doc.paragraphs:
                parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    parts.append(row_text)
            text = "\n".join(parts)
        else:
            # fallback: try pytesseract on file as image
            try:
                img = Image.open(filepath)
                ocr_res = ocr_image_with_confidence(img)
                text = ocr_res['text']
                avg_conf = ocr_res['avg_confidence']
                source = 'image'
            except Exception:
                text = ""
    except Exception as e:
        print(f"Error extracting text from {filepath}: {e}")
    return {'text': text or "", 'source': source or "unknown", 'avg_conf': avg_conf}

# -------------------------
# Aadhaar hashing + lookup
# -------------------------
def hash_aadhaar(aadhaar: str) -> str:
    """
    Returns a hex SHA-256 hash of aadhaar + salt. Requires AADHAAR_SALT env var.
    """
    salt = os.getenv('AADHAAR_SALT')
    if not salt:
        raise EnvironmentError("AADHAAR_SALT environment variable not set. Set it before running.")
    cleaned = re.sub(r'\D', '', aadhaar).strip()
    return hashlib.sha256((cleaned + salt).encode('utf-8')).hexdigest()

def find_user_by_aadhaar(aadhaar: str):
    """
    Returns user_id if an existing user has the same aadhaar hash, else None.
    """
    try:
        h = hash_aadhaar(aadhaar)
    except EnvironmentError:
        return None
    conn = get_db_conn()
    cur = conn.cursor()
    cur.execute("SELECT user_id FROM users WHERE aadhaar_hash = ?", (h,))
    row = cur.fetchone()
    conn.close()
    return row['user_id'] if row else None

# -------------------------
# AI / RAG helpers (Gemini prompts, same as before)
# -------------------------
def analyze_form_fields_with_rag(template_path):
    """Analyze DOCX to get candidate fields and ask model to consolidate & output JSON list."""
    if not model:
        flash("AI Model is not configured. Please set the GEMINI_API_KEY.", "danger")
        return []

    try:
        doc = Document(template_path)
        raw_fields = []
        # Heuristics: table cells
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text and (':' in text or len(text.split()) < 5) and cell_idx + 1 < len(row.cells):
                        raw_fields.append({
                            "field_id": f"table_{table_idx}_row_{row_idx}_cell_{cell_idx+1}",
                            "label": text.replace(':', '').strip()
                        })
        # Heuristics: paragraph labels
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if ':' in text and len(text.split(':')[-1].strip()) < 10:
                raw_fields.append({
                    "field_id": f"para_{para_idx}",
                    "label": text.split(':')[0].strip()
                })

        if not raw_fields:
            return []

        prompt = f"""
You are an AI expert at analyzing Indian government application forms.
Based on the following list of field labels extracted from a form, please process them.

Extracted Labels:
{json.dumps(raw_fields, indent=2)}

Your tasks are to:
1. Consolidate semantically duplicate fields (e.g., "Applicant Name" and "Full Name").
2. Generate a standardized, snake_case `field_name` for each unique field (e.g., 'father_name', 'date_of_birth').
3. Determine the most appropriate HTML input `field_type` (e.g., 'text', 'date', 'number', 'email', 'tel', 'select').
4. Assign a `priority` score from 1 (least important) to 10 (most important).
5. Keep the original `field_id` for mapping. If consolidating, choose the most appropriate `field_id`.

Return ONLY a valid JSON array of objects, where each object represents a unique field.
Example format:
[
  {
    "field_id": "table_0_row_1_cell_1",
    "field_name": "applicant_name",
    "label": "Applicant Name",
    "field_type": "text",
    "priority": 10
  }
]
JSON Output:
"""
        response = model.generate_content(prompt)
        json_string = response.text.strip().replace('```json', '').replace('```', '')
        enhanced_fields = json.loads(json_string)
        return enhanced_fields
    except Exception as e:
        print(f"Error analyzing form fields with RAG: {e}")
        flash(f"AI could not analyze the form. Error: {e}", "warning")
        return []

def get_structured_data_with_rag(documents_text, required_fields):
    """Use Gemini to extract values for required fields from the combined document text."""
    if not model or not documents_text:
        return {}

    fields_json = json.dumps([{"field_name": f["field_name"], "label": f["label"]} for f in required_fields], indent=2)
    prompt = f"""
You are an AI assistant specialized in extracting data from Indian KYC and land documents.
Based on the **Required Fields** list below, extract the corresponding information from the **Document Text**.

**Required Fields (JSON format):**
{fields_json}

**Document Text:**
---
{documents_text}
---

**Instructions:**
1. Carefully map the information from the text to the `field_name` in the required fields list.
2. Pay close attention to details like names, dates (format as YYYY-MM-DD), and multi-digit numbers (Aadhaar, Account numbers).
3. If a piece of information for a required field is not found, use an empty string `""` as its value.
4. Return ONLY a single, valid JSON object where keys are the `field_name`s from the list.

JSON Output:
"""
    try:
        response = model.generate_content(prompt)
        json_string = response.text.strip().replace('```json', '').replace('```', '')
        return json.loads(json_string)
    except Exception as e:
        print(f"Error calling Gemini API or parsing JSON: {e}")
        return {}

def fill_form_template_precise(template_path, form_data, output_name):
    """Fill the DOCX template based on field_id placements."""
    try:
        doc = Document(template_path)
        filled_any = False
        for field_id, value in form_data.items():
            if not value:
                continue
            try:
                parts = field_id.split('_')
                if field_id.startswith('table_'):
                    table_idx, row_idx, cell_idx = int(parts[1]), int(parts[3]), int(parts[5])
                    if table_idx < len(doc.tables) and row_idx < len(doc.tables[table_idx].rows) and cell_idx < len(doc.tables[table_idx].rows[row_idx].cells):
                        target_cell = doc.tables[table_idx].rows[row_idx].cells[cell_idx]
                        target_cell.text = value
                        filled_any = True
                elif field_id.startswith('para_'):
                    para_idx = int(parts[1])
                    if para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[para_idx]
                        if ':' in para.text:
                            para.text = para.text.split(':')[0] + ': ' + value
                        else:
                            para.text = value
                        filled_any = True
            except (ValueError, IndexError) as e:
                print(f"Could not parse or find position for field_id {field_id}: {e}")
                continue

        if not filled_any:
            print("Warning: No fields were filled in the document.")
            return None

        timestamp = str(int(time.time()))
        output_filename = f"{output_name}-filled-{timestamp}.docx"
        output_path = os.path.join(app.config['GENERATED_FOLDER'], output_filename)
        doc.save(output_path)
        return output_filename
    except Exception as e:
        print(f"Error filling form template: {e}")
        return None

# -------------------------
# Routes: ingest & auto_fill_user
# -------------------------
@app.route('/ingest', methods=['POST'])
def ingest_documents():
    """
    Ingest uploaded docs for a user. Returns JSON with user_id UUID.
    If client sends user_id it will be used; otherwise a new UUID is generated.
    If an Aadhaar-like number is found in the text it is hashed and stored (not the raw Aadhaar).
    """
    scheme_id = request.form.get('scheme')
    provided_user_id = request.form.get('user_id', '').strip() or None
    files = request.files.getlist('documents')

    if not scheme_id or scheme_id not in SCHEMES:
        return jsonify({'error': 'Please provide a valid scheme id.'}), 400

    if not files or all(f.filename == '' for f in files):
        return jsonify({'error': 'Please upload one or more documents.'}), 400

    saved_filenames = []
    per_file_texts = []  # list of (filename, extracted_text, avg_conf, source)
    inferred_aadhaar = None

    # Save files & extract text
    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            saved_filenames.append(filename)

            res = extract_text_from_file(filepath)
            text = res.get('text', '') or ''
            avg_conf = res.get('avg_conf', None)
            source = res.get('source', None)
            per_file_texts.append((filename, text, avg_conf, source))

            if not inferred_aadhaar:
                a = find_aadhaar_in_text(text)
                if a:
                    inferred_aadhaar = a

    # Determine user_id (client-provided preferred, else new UUID)
    user_id = provided_user_id or str(uuid.uuid4())

    # Hash detected aadhaar if present
    aadhaar_hash = None
    if inferred_aadhaar:
        try:
            aadhaar_hash = hash_aadhaar(inferred_aadhaar)
        except EnvironmentError:
            # AADHAAR_SALT not set; aadhaar_hash stays None
            aadhaar_hash = None

    # Save user (or update if exists)
    save_user_if_new(user_id, aadhaar_hash=aadhaar_hash)

    # Save documents into DB. Chunk text and store metadata.
    for filename, text, avg_conf, source in per_file_texts:
        if not text.strip():
            save_document_record(user_id, filename, scheme_id, "", doc_type=source, metadata={'ocr_conf': avg_conf or 0}, chunk_index=-1)
            continue

        chunks = chunk_text_simple(text, words_per_chunk=400)
        if not chunks:
            save_document_record(user_id, filename, scheme_id, text, doc_type=source, metadata={'ocr_conf': avg_conf or 0}, chunk_index=-1)
            continue

        for idx, chunk in enumerate(chunks):
            meta = {'ocr_conf': avg_conf or 0, 'orig_filename': filename}
            save_document_record(user_id, filename, scheme_id, chunk, doc_type=source, metadata=meta, chunk_index=idx)

    return jsonify({
        'message': 'Documents ingested successfully.',
        'user_id': user_id,
        'files_saved': saved_filenames
    }), 200

@app.route('/auto_fill_user', methods=['POST'])
def auto_fill_user():
    """
    Unified endpoint to auto-fill forms for a user.
    Accepts form-data or JSON. Key parameters:
      - user_id OR aadhaar
      - scheme (optional) OR form_file (optional)
      - fields (optional JSON array) -> client-provided form fields (for HTML/extension)
      - output_type: 'pdf' (default), 'html', or 'json'
    Behavior:
      - pdf: fill DOCX template (server scheme or uploaded form_file) and return download URL JSON.
      - html: render HTML form for browser (if Accept: text/html) OR return JSON with fields + html_preview.
      - json: return structured mapped values as JSON.
    """
    # Accept both JSON and form-encoded requests
    data = request.get_json(silent=True) or {}
    # Form fallback
    user_id = (data.get('user_id') or request.form.get('user_id') or '').strip()
    aadhaar_input = (data.get('aadhaar') or request.form.get('aadhaar') or '').strip()
    scheme_id = (data.get('scheme') or request.form.get('scheme') or '').strip()
    output_type = (data.get('output_type') or request.form.get('output_type') or 'pdf').strip().lower()

    # fields: prefer JSON body 'fields', else form field 'fields' as JSON string
    fields_payload = data.get('fields') or request.form.get('fields')
    if isinstance(fields_payload, str):
        try:
            fields_payload = json.loads(fields_payload)
        except Exception:
            fields_payload = None

    # uploaded form file (optional) for ad-hoc pdf/docx
    form_file = request.files.get('form_file')

    # Validate identification
    if not user_id and not aadhaar_input:
        return jsonify({"success": False, "error": "Provide user_id or aadhaar"}), 400

    if not user_id and aadhaar_input:
        found = find_user_by_aadhaar(aadhaar_input)
        if not found:
            return jsonify({"success": False, "error": "No user found for provided Aadhaar"}), 404
        user_id = found

    # Load user's stored docs
    docs = get_documents_by_user(user_id)
    if not docs:
        return jsonify({"success": False, "error": "No documents found for this user_id"}), 404
    combined_text = "\n\n".join([d['text'] for d in docs if d.get('text')])
    if not combined_text.strip():
        return jsonify({"success": False, "error": "Stored documents contain no usable text"}), 400

    # Determine template_path if needed (pdf flow)
    template_path = None
    if form_file:
        # save uploaded form file temporarily to uploads and use it as template
        fname = secure_filename(form_file.filename)
        fp = os.path.join(app.config['UPLOAD_FOLDER'], fname)
        form_file.save(fp)
        template_path = fp
    elif scheme_id:
        scheme_info = SCHEMES.get(scheme_id)
        if scheme_info:
            template_path = os.path.join(app.config['TEMPLATE_FOLDER'], scheme_info['template_file'])
            if not os.path.exists(template_path):
                template_path = None

    # Build required_fields list (for RAG)
    # Priority: client-provided fields_payload -> if not present, template-derived fields (if template_path)
    required_fields = []
    field_client_map = {}  # maps client form_field_id -> desired field_name (to map returned values)
    if fields_payload and isinstance(fields_payload, list):
        # Client fields expected: [{"field_id":"f1","label":"Full name","field_type":"text"}, ...]
        for f in fields_payload:
            label = f.get('label') or f.get('name') or f.get('field_id') or ''
            field_name = generate_field_name_from_label(label)
            required_fields.append({"field_name": field_name, "label": label})
            # map back later using client id
            field_client_map[field_name] = f.get('field_id') or label
    else:
        # No client-provided fields: try to derive from template
        if template_path:
            required_fields = analyze_form_fields_with_rag(template_path)
            # analyze_form_fields_with_rag returns objects with field_name and field_id (doc positions).
            # For mapping back to client targets, we'll use the returned 'field_name' and keep doc field_ids as keys.
        else:
            # if neither client fields nor template available, can't proceed
            return jsonify({"success": False, "error": "No form fields provided and no template available to analyze."}), 400

    # Call RAG to extract structured data
    extracted_data = get_structured_data_with_rag(combined_text, required_fields)
    if extracted_data is None:
        return jsonify({"success": False, "error": "AI failed to extract structured data"}), 500

    # Map RAG extracted data back to the output mapping expected by client or by template
    mapped_fields = {}
    # Case A: client provided fields -> map field_name -> client field_id
    if fields_payload and isinstance(fields_payload, list):
        for rf in required_fields:
            fname = rf.get('field_name')
            client_fid = field_client_map.get(fname) or fname
            value = extracted_data.get(fname, "") if isinstance(extracted_data, dict) else ""
            mapped_fields[client_fid] = {"value": str(value), "confidence": None, "source": "rag"}
    else:
        # No client fields; we derived required_fields from template (which has 'field_id' positions)
        # required_fields likely look like: [{ "field_id":"para_3", "field_name":"applicant_name", ... }, ...]
        # If analyze returned these richer objects, use field['field_id'] as key to map to extracted_data[field_name]
        for f in required_fields:
            field_id = f.get('field_id') or f.get('field_name')
            fname = f.get('field_name')
            value = extracted_data.get(fname, "") if isinstance(extracted_data, dict) else ""
            mapped_fields[field_id] = {"value": str(value), "confidence": None, "source": "rag"}

    # Now branch for output_type
    try:
        if output_type == 'pdf':
            # Must have a template_path
            if not template_path:
                return jsonify({"success": False, "error": "No template provided for PDF flow (provide scheme or upload form_file)."}), 400

            # fill_form_template_precise expects mapping doc_field_id -> value
            # If we used client fields, we don't have doc_field_ids; attempt to map using field_name -> doc field id if available
            form_fill_map = {}
            if fields_payload and isinstance(fields_payload, list) and template_path:
                # attempt to use analyze_form_fields_with_rag(template_path) to map field_names -> doc field_id
                tpl_fields = analyze_form_fields_with_rag(template_path)
                name_to_id = {f.get('field_name'): f.get('field_id') for f in tpl_fields if f.get('field_name') and f.get('field_id')}
                for fname, client_fid in field_client_map.items():
                    doc_id = name_to_id.get(fname)
                    if doc_id and mapped_fields.get(client_fid):
                        form_fill_map[doc_id] = mapped_fields[client_fid]['value']
            else:
                # template-derived mapping stored already in mapped_fields keyed by doc field ids
                form_fill_map = {k: v['value'] for k, v in mapped_fields.items()}

            if not form_fill_map:
                return jsonify({"success": False, "error": "Could not map extracted values to template fields for PDF fill."}), 400

            filled_filename = fill_form_template_precise(template_path, form_fill_map, (scheme_id or "form") + "_" + user_id)
            if filled_filename:
                download_url = url_for('download_page', filename=filled_filename, _external=True)
                return jsonify({"success": True, "mode":"pdf", "user_id": user_id, "filled_form": download_url})
            else:
                return jsonify({"success": False, "error": "Failed to fill template file."}), 500

        elif output_type == 'html':
            # If browser client: render HTML with mapped fields; else return JSON + html_preview
            accept = request.headers.get('Accept','')
            # Prepare fields for template (key -> value)
            html_fields = {k: v['value'] for k,v in mapped_fields.items()}

            # If client gave client-style field ids, these keys map directly to form inputs on page
            if 'text/html' in accept:
                # Render the actual HTML form for browser interaction
                return render_template('filled_form.html', fields=html_fields, user_id=user_id, scheme_id=scheme_id)
            else:
                # Return JSON, include a small html preview so clients can show a quick UI preview
                html_preview = render_template('filled_form.html', fields=html_fields, user_id=user_id, scheme_id=scheme_id)
                return jsonify({
                    "success": True,
                    "mode": "html",
                    "user_id": user_id,
                    "scheme": scheme_id,
                    "fields": html_fields,
                    "html_preview": html_preview  # full html (can be large) - client can choose to display
                })

        elif output_type == 'json':
            return jsonify({
                "success": True,
                "mode": "json",
                "user_id": user_id,
                "scheme": scheme_id,
                "mapped_fields": mapped_fields,
                "diagnostics": {"doc_count": len(docs)}
            })

        else:
            return jsonify({"success": False, "error": "Invalid output_type"}), 400

    except Exception as e:
        return jsonify({"success": False, "error": f"Unexpected error: {str(e)}"}), 500

# -------------------------
# Existing routes: front-end helpers
# -------------------------
@app.route('/')
def index():
    return render_template('index.html', schemes=SCHEMES)

@app.route('/manual', methods=['GET', 'POST'])
def manual_fill():
    if not model:
        flash("AI Model is not configured due to missing API key. Please contact the administrator.", "danger")
        return redirect(url_for('index'))

    if request.method == 'POST':
        step = request.form.get('step')
        if step == '1':
            form_file = request.files.get('form_template')
            support_docs = request.files.getlist('support_documents')

            if not form_file or not allowed_file(form_file.filename) or not form_file.filename.endswith('.docx'):
                flash('Please upload a valid DOCX form template.', 'danger')
                return redirect(request.url)

            try:
                form_filename = secure_filename(form_file.filename)
                form_path = os.path.join(app.config['UPLOAD_FOLDER'], form_filename)
                form_file.save(form_path)

                required_fields = analyze_form_fields_with_rag(form_path)
                if not required_fields:
                    flash('Could not analyze the form. Ensure it contains identifiable fields (like "Name:", "Address:", etc.).', 'danger')
                    return redirect(request.url)

                extracted_data = {}
                if support_docs and any(f.filename for f in support_docs):
                    combined_text = ""
                    for doc in support_docs:
                        if doc and allowed_file(doc.filename):
                            doc_filename = secure_filename(doc.filename)
                            filepath = os.path.join(app.config['UPLOAD_FOLDER'], doc_filename)
                            doc.save(filepath)
                            combined_text += extract_text_from_file(filepath)['text'] + "\n\n"
                    
                    if combined_text.strip():
                        extracted_data = get_structured_data_with_rag(combined_text, required_fields)

                if extracted_data:
                    for field in required_fields:
                        if field['field_name'] in extracted_data and extracted_data[field['field_name']]:
                            field['pre_filled_value'] = extracted_data[field['field_name']]

                return render_template('manual_fill_step2.html', 
                                     fields=required_fields, 
                                     form_filename=form_filename,
                                     has_extracted_data=bool(extracted_data))
            except Exception as e:
                flash(f'Error processing your request: {e}', 'danger')
                return redirect(request.url)
        
        elif step == '2':
            form_filename = request.form.get('form_filename')
            form_path = os.path.join(app.config['UPLOAD_FOLDER'], form_filename)

            if not form_filename or not os.path.exists(form_path):
                flash('Form template not found. Please start over.', 'danger')
                return redirect(url_for('manual_fill'))
            
            form_data = {key: value for key, value in request.form.items() if key not in ['step', 'form_filename']}
            filled_form_filename = fill_form_template_precise(form_path, form_data, "manual-form")

            if filled_form_filename:
                flash('Form filled successfully! Please review the downloaded document.', 'success')
                return redirect(url_for('download_page', filename=filled_form_filename))
            else:
                flash('An error occurred while filling the form. No data was entered.', 'danger')
                return redirect(url_for('manual_fill'))

    return render_template('manual_fill.html')

@app.route('/get_scheme_info/<scheme_id>')
def get_scheme_info(scheme_id):
    scheme = SCHEMES.get(scheme_id)
    if scheme:
        return jsonify(scheme)
    return jsonify({'error': 'Scheme not found'}), 404

@app.route('/download_page/<filename>')
def download_page(filename):
    return render_template('download.html', filename=filename)

@app.route('/generated/<filename>')
def download_form(filename):
    return send_from_directory(app.config['GENERATED_FOLDER'], filename, as_attachment=True)

# -------------------------
# Startup
# -------------------------
if __name__ == '__main__':
    # Create necessary folders on startup
    for folder in [UPLOAD_FOLDER, GENERATED_FOLDER, TEMPLATE_FOLDER]:
        os.makedirs(folder, exist_ok=True)

    # Init DB
    init_db()

    app.run(debug=True)
    CORS(app, supports_credentials=True)
